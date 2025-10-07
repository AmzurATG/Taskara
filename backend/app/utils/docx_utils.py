from docx import Document
import requests
import tempfile
import os
from typing import List, Dict, Any
import re


class DOCXExtractor:
    @staticmethod
    def _download_file_from_url(url: str) -> str:
        """Download file from URL and return temporary file path."""
        try:
            # Clean URL - remove any trailing characters
            url = url.rstrip('?')
            
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            
            # Create temporary file
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            temp_file.write(response.content)
            temp_file.close()
            
            return temp_file.name
        except Exception as e:
            raise Exception(f"Error downloading file from URL: {str(e)}")
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text from DOCX file (supports both local paths and URLs)."""
        temp_file_path = None
        try:
            # Check if it's a URL
            if file_path.startswith('http'):
                temp_file_path = DOCXExtractor._download_file_from_url(file_path)
                actual_file_path = temp_file_path
            else:
                actual_file_path = file_path
            
            doc = Document(actual_file_path)
            text = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            return text.strip()
        
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")
        
        finally:
            # Clean up temporary file if it was created
            if temp_file_path and os.path.exists(temp_file_path):
                try:
                    os.unlink(temp_file_path)
                except Exception:
                    pass
    
    @staticmethod
    def extract_structured_content(file_path: str) -> Dict[str, Any]:
        """Extract structured content from DOCX file (supports both local paths and URLs)."""
        temp_file_path = None
        try:
            # Check if it's a URL
            if file_path.startswith('http'):
                temp_file_path = DOCXExtractor._download_file_from_url(file_path)
                actual_file_path = temp_file_path
            else:
                actual_file_path = file_path
            
            doc = Document(actual_file_path)
            content = {
                "paragraphs": [],
                "tables": [],
                "headings": [],
                "lists": []
            }
            
            # Extract paragraphs with styles
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    para_info = {
                        "text": paragraph.text.strip(),
                        "style": paragraph.style.name if paragraph.style else "Normal",
                        "is_heading": paragraph.style.name.startswith('Heading') if paragraph.style else False
                    }
                    content["paragraphs"].append(para_info)
                    
                    # Separate headings
                    if para_info["is_heading"]:
                        content["headings"].append({
                            "text": para_info["text"],
                            "level": para_info["style"]
                        })
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                
                content["tables"].append({
                    "index": table_idx,
                    "data": table_data
                })
            
            return content
        
        except Exception as e:
            raise Exception(f"Error extracting structured content from DOCX: {str(e)}")
        
        finally:
            # Clean up temporary file if it was created
            if temp_file_path and os.path.exists(temp_file_path):
                try:
                    os.unlink(temp_file_path)
                except Exception:
                    pass
    
    @staticmethod
    def extract_by_sections(file_path: str) -> List[Dict[str, Any]]:
        """Extract content organized by sections/headings (supports both local paths and URLs)."""
        try:
            structured_content = DOCXExtractor.extract_structured_content(file_path)
            sections = []
            current_section = None
            
            for para in structured_content["paragraphs"]:
                if para["is_heading"]:
                    # Start new section
                    if current_section:
                        sections.append(current_section)
                    
                    current_section = {
                        "heading": para["text"],
                        "level": para["style"],
                        "content": []
                    }
                else:
                    # Add to current section or create default section
                    if current_section is None:
                        current_section = {
                            "heading": "Introduction",
                            "level": "Default",
                            "content": []
                        }
                    
                    current_section["content"].append(para["text"])
            
            # Add last section
            if current_section:
                sections.append(current_section)
            
            return sections
        
        except Exception as e:
            raise Exception(f"Error extracting sections from DOCX: {str(e)}")
